#!/usr/bin/env python3
"""
Generate QLTB document from JSON OCR data and Word template.
Updated to handle schema from bbbgtaq7 repository.
"""

import argparse
import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
import base64
import qrcode
from PIL import Image
from typing import Dict, Any

def generate_qr_code(data: str, size: int = 5) -> str:
    """Generate QR code as base64 PNG."""
    qr = qrcode.QRCode(version=1, box_size=10, border=5)
    qr.add_data(data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img = img.resize((size * 10, size * 10))
    from io import BytesIO
    buffered = BytesIO()
    img.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode()

def replace_placeholders(doc, data: Dict[str, Any]) -> None:
    """Replace placeholders in Word document with data."""
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = f'{{{{{key}}}}}'
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

def main():
    parser = argparse.ArgumentParser(description='Generate QLTB document from JSON OCR data')
    parser.add_argument('--input', required=True, help='Input JSON file path')
    parser.add_argument('--template', required=True, help='Template Word file path (.docx)')
    parser.add_argument('--output', required=True, help='Output Word file path')
    args = parser.parse_args()

    with open(args.input, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Generate QR code if serial number exists
    if 'serial' in data and 'model' in data:
        qr_data = f"asset:{data.get('asset_tag', '')}|serial:{data['serial']}|model:{data['model']}"
        qr_base64 = generate_qr_code(qr_data)
        data['qr_code'] = qr_base64

    # Load and fill template
    doc = Document(args.template)
    replace_placeholders(doc, data)
    doc.save(args.output)
    print(f"Document generated: {args.output}")

if __name__ == '__main__':
    main()