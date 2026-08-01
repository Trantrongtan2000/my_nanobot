# -*- coding: utf-8 -*-
"""Second pass: extract DOCX attachments + Q7 package map; enrich concept pages."""
from __future__ import annotations

import json
import re
import unicodedata
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document

ROOT = Path.home() / "Downloads" / "New folder (3)"
WS = Path.home() / ".nanobot" / "workspace"
WIKI = WS / "wiki"
RAW = WIKI / "raw" / "ksnk"
EXTRACT = RAW / "extracts" / "docx"
CONCEPTS = WIKI / "concepts"
SYN = WIKI / "synthesis"
TODAY = date.today().isoformat()
MAX_CHARS = 12000


def slugify(s: str) -> str:
    s = s.lower().strip().replace("đ", "d").replace("Đ", "d")
    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    s = re.sub(r"[^a-z0-9]+", "_", s)
    return re.sub(r"_+", "_", s).strip("_")[:80]


def site_of(rel: Path) -> str:
    p0 = rel.parts[0] if rel.parts else ""
    if "7" in p0 and ("Qu" in p0 or "qu" in p0.lower() or "ận" in p0 or "an" in p0):
        return "q7"
    if "Tân" in p0 or "Tan" in p0 or "Bình" in p0:
        return "tb"
    # fallback by path string
    s = str(rel)
    if "Quận 7" in s or "Quan 7" in s:
        return "q7"
    if "Tân Bình" in s or "Tan Binh" in s:
        return "tb"
    return "unk"


def doc_code(name: str) -> str | None:
    n = name
    patterns = [
        r"(TA[25]\.KSNK\.(?:QT|QĐ|QD|HD)\.\d+[A-Z]?)",
        r"(KSNK\.(?:QT|QĐ|QD|HD)\.\d+[A-Z]?)",
        r"(CS\.KSNK\.\d+)",
        r"(QTKT\.KSNK\.\d+)",
        r"(QT\.KSNK\.\d+)",
    ]
    for pat in patterns:
        m = re.search(pat, n, re.I)
        if m:
            code = m.group(1)
            # normalize QD -> QĐ
            code = re.sub(r"\.QD\.", ".QĐ.", code, flags=re.I)
            # preserve case pattern TA5.KSNK.QT.02
            parts = code.split(".")
            if len(parts) >= 4:
                parts[0] = parts[0].upper()
                parts[1] = parts[1].upper()
                parts[2] = parts[2].upper().replace("QD", "QĐ")
                if parts[2] == "QĐ":
                    pass
                code = ".".join(parts[:3] + [parts[3].upper()])
            return code
    return None


def norm_key(code: str) -> str:
    c = code.upper()
    c = c.replace("QĐ", "QD")
    c = re.sub(r"^TA[25]\.", "", c)
    c = c.replace("QD", "QĐ")
    # QT.KSNK.09 -> keep
    if c.startswith("QT.KSNK"):
        return c
    if c.startswith("QTKT.KSNK"):
        return c
    if c.startswith("CS.KSNK"):
        return c
    return c


def extract_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
        paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                line = " | ".join(cells)
                if line.strip(" |"):
                    paras.append(line)
        text = "\n".join(paras)
        if len(text) > MAX_CHARS:
            text = text[:MAX_CHARS] + "\n\n[...truncated...]"
        return text
    except Exception as e:
        return f"[extract_error] {e}"


def att_kind(name: str) -> str:
    m = re.match(r"^(BK|BM|PL|QA)(\d*)", name, re.I)
    if m:
        return m.group(1).upper()
    return "DOCX"


def main():
    EXTRACT.mkdir(parents=True, exist_ok=True)
    docxs = [
        p
        for p in ROOT.rglob("*")
        if p.is_file()
        and p.suffix.lower() in {".docx", ".doc"}
        and not p.name.startswith("~$")
    ]
    print(f"docx_count={len(docxs)}", flush=True)

    # Only process .docx (python-docx); skip .doc
    docxs = [p for p in docxs if p.suffix.lower() == ".docx"]

    by_key_site: dict[tuple[str, str], list[dict]] = defaultdict(list)
    all_att = []

    for p in sorted(docxs, key=lambda x: str(x).lower()):
        rel = p.relative_to(ROOT)
        site = site_of(rel)
        code = doc_code(p.name) or doc_code(str(rel))
        if not code:
            # try parent folder
            for part in rel.parts:
                code = doc_code(part)
                if code:
                    break
        if not code:
            code = "UNCODED"
        key = norm_key(code) if code != "UNCODED" else "UNCODED"
        text = extract_docx(p)
        kind = att_kind(p.name)
        ex_name = f"{site}_{slugify(code)}_{slugify(p.stem)[:50]}.md"
        ex_path = EXTRACT / ex_name
        body = "\n".join(
            [
                "---",
                "type: source",
                f'title: "{p.name}"',
                "status: draft",
                f"updated: {TODAY}",
                f'sources: ["{rel.as_posix()}"]',
                "tags: [ksnk, docx, attachment]",
                "---",
                "",
                f"# {p.name}",
                "",
                f"- Site: {site}",
                f"- Code: {code}",
                f"- Kind: {kind}",
                f"- Path: `{p}`",
                f"- Relative: `{rel.as_posix()}`",
                "",
                "## Text",
                "",
                "```",
                text if text.strip() else "[empty]",
                "```",
                "",
            ]
        )
        ex_path.write_text(body, encoding="utf-8")
        rec = {
            "site": site,
            "code": code,
            "key": key,
            "kind": kind,
            "name": p.name,
            "rel": rel.as_posix(),
            "extract": f"raw/ksnk/extracts/docx/{ex_name}",
            "chars": len(text),
            "preview": re.sub(r"\s+", " ", text)[:400],
            "has_text": len(text) > 80 and not text.startswith("[extract_error]"),
        }
        by_key_site[(key, site)].append(rec)
        all_att.append(rec)
        msg = f"docx {site} {code} {kind} chars={len(text)}"
        print(msg.encode("ascii", "replace").decode(), flush=True)

    (RAW / "docx_index.json").write_text(
        json.dumps(all_att, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # Q7 package map from folder structure under QTVH
    q7root = None
    for d in ROOT.iterdir():
        if d.is_dir() and "7" in d.name:
            q7root = d
            break
    packages = []
    if q7root:
        ksnk = None
        for d in q7root.rglob("*"):
            if d.is_dir() and d.name.upper().startswith("QTVH"):
                # parent is KSNK folder
                for child in sorted(d.iterdir()):
                    if child.is_dir():
                        code = doc_code(child.name) or ""
                        files = [
                            f
                            for f in child.rglob("*")
                            if f.is_file() and not f.name.startswith("~$")
                        ]
                        packages.append(
                            {
                                "folder": child.name,
                                "code": code,
                                "rel": child.relative_to(ROOT).as_posix(),
                                "files": [
                                    {
                                        "name": f.name,
                                        "kb": f.stat().st_size // 1024,
                                        "ext": f.suffix.lower(),
                                        "kind": att_kind(f.name)
                                        if f.suffix.lower() == ".docx"
                                        else f.suffix.lower().lstrip("."),
                                    }
                                    for f in sorted(files, key=lambda x: x.name.lower())
                                ],
                            }
                        )
                    elif child.is_file() and child.suffix.lower() == ".pdf":
                        packages.append(
                            {
                                "folder": child.stem,
                                "code": doc_code(child.name) or "",
                                "rel": child.relative_to(ROOT).as_posix(),
                                "files": [
                                    {
                                        "name": child.name,
                                        "kb": child.stat().st_size // 1024,
                                        "ext": ".pdf",
                                        "kind": "pdf",
                                    }
                                ],
                            }
                        )
                break  # first QTVH under q7

    # Write Q7 package synthesis
    pkg_path = SYN / "ksnk_q7_goi_quy_trinh.md"
    plines = [
        "---",
        "type: synthesis",
        'title: "KSNK Quận 7 — Gói quy trình QTVH (map thư mục)"',
        "status: draft",
        "sources:",
        '  - "raw/ksnk/catalog.md"',
        f'  - "source-root:{ROOT.as_posix()}"',
        f"updated: {TODAY}",
        "tags: [ksnk, q7, package-map]",
        "refs: []",
        "---",
        "",
        "# KSNK Quận 7 — Gói quy trình (QTVH)",
        "",
        "Map theo cấu trúc thư mục nguồn. PDF chính là **scan** (không text-layer); nội dung chi tiết lấy từ **BK/BM/PL .docx** khi có.",
        "",
        f"- Nguồn: `{ROOT}`",
        f"- Số gói QTVH: **{len(packages)}**",
        "- Hub chung: [[synthesis/ksnk_quy_trinh_hub]]",
        "",
    ]
    for pkg in sorted(packages, key=lambda x: x.get("code") or x["folder"]):
        plines.append(f"## {pkg['code'] or '—'} — {pkg['folder']}")
        plines.append("")
        plines.append(f"- Path: `{pkg['rel']}`")
        key = norm_key(pkg["code"]) if pkg["code"] else ""
        # link concept if exists
        if key:
            # find concept file
            hits = list(CONCEPTS.glob(f"ksnk_{slugify(key)}_*.md"))
            if not hits:
                hits = list(CONCEPTS.glob(f"*{slugify(key)}*.md"))
            if hits:
                plines.append(f"- Concept: [[concepts/{hits[0].stem}]]")
        plines.append("- Thành phần:")
        for f in pkg["files"]:
            plines.append(f"  - `{f['kind']}` {f['name']} ({f['kb']} KB)")
        # attach docx previews for this code
        if pkg["code"]:
            atts = by_key_site.get((norm_key(pkg["code"]), "q7"), [])
            if atts:
                plines.append("- Nội dung đính kèm (trích DOCX):")
                for a in atts:
                    plines.append(f"  - **{a['kind']}** `{a['name']}` ({a['chars']} ký tự)")
                    if a["preview"]:
                        plines.append(f"    - Preview: {a['preview'][:300]}")
                    plines.append(f"    - Extract: `{a['extract']}`")
        plines.append("")

    pkg_path.write_text("\n".join(plines), encoding="utf-8")
    print(f"pkg_map -> {pkg_path}", flush=True)

    # Enrich existing concept pages with DOCX sections
    # Build map key -> concept file
    concept_files = list(CONCEPTS.glob("ksnk_*.md"))
    enriched = 0
    for cf in concept_files:
        text = cf.read_text(encoding="utf-8")
        # find group key from page
        m = re.search(r"Nhóm mã:\s*`([^`]+)`", text)
        if not m:
            continue
        key = m.group(1)
        atts = []
        for (k, site), items in by_key_site.items():
            if k == key or k.replace("QĐ", "QD") == key.replace("QĐ", "QD"):
                atts.extend(items)
        if not atts:
            continue
        # remove old attachments section if re-run
        if "## Đính kèm BK/BM/PL" in text:
            text = text.split("## Đính kèm BK/BM/PL")[0].rstrip() + "\n"
        section = ["", "## Đính kèm BK/BM/PL (text từ DOCX)", ""]
        for a in sorted(atts, key=lambda x: (x["site"], x["kind"], x["name"])):
            flag = "ok" if a["has_text"] else "weak"
            section.append(
                f"### {a['site'].upper()} · {a['kind']} · {a['name']} ({flag})"
            )
            section.append("")
            section.append(f"- Nguồn: `{a['rel']}`")
            section.append(f"- Extract: `{a['extract']}`")
            section.append("")
            if a["has_text"]:
                section.append("```")
                # full-ish preview up to 2500
                body = Path(WS / "wiki" / a["extract"]).read_text(encoding="utf-8")
                if "```" in body:
                    chunk = body.split("```", 2)[1]
                    if chunk.startswith("\n"):
                        chunk = chunk[1:]
                else:
                    chunk = a["preview"]
                section.append(chunk[:2500])
                section.append("```")
            section.append("")
        cf.write_text(text.rstrip() + "\n" + "\n".join(section) + "\n", encoding="utf-8")
        enriched += 1
        print(f"enriched {cf.name} atts={len(atts)}", flush=True)

    # Update hub with Q7 package link + docx stats
    hub = SYN / "ksnk_quy_trinh_hub.md"
    hub_t = hub.read_text(encoding="utf-8")
    if "ksnk_q7_goi_quy_trinh" not in hub_t:
        insert = (
            "\n## Map gói Quận 7\n\n"
            f"- [[synthesis/ksnk_q7_goi_quy_trinh]] — {len(packages)} gói QTVH + thành phần BK/BM/PL\n"
            f"- DOCX attachments extracted: **{len(all_att)}** (text-ok: {sum(1 for a in all_att if a['has_text'])})\n"
            f"- Index DOCX: `raw/ksnk/docx_index.json`\n"
        )
        # insert after Phạm vi nguồn section
        hub_t = hub_t.replace(
            "## Cấu trúc mã tài liệu",
            insert + "\n## Cấu trúc mã tài liệu",
        )
        hub.write_text(hub_t, encoding="utf-8")

    # index
    idx_path = WIKI / "index.md"
    idx = idx_path.read_text(encoding="utf-8")
    if "ksnk_q7_goi_quy_trinh" not in idx:
        idx = idx.replace(
            "- [[synthesis/ksnk_quy_trinh_hub]]",
            "- [[synthesis/ksnk_quy_trinh_hub]]\n"
            f"- [[synthesis/ksnk_q7_goi_quy_trinh]] — Map gói QTVH Quận 7 + trích BK/BM/PL ({TODAY})",
        )
        idx_path.write_text(idx, encoding="utf-8")

    log_path = WIKI / "log.md"
    log = log_path.read_text(encoding="utf-8")
    log += (
        f"\n## {TODAY} — enrich KSNK DOCX\n"
        f"- DOCX extracted: {len(all_att)} → `raw/ksnk/extracts/docx/`\n"
        f"- Concept pages enriched: {enriched}\n"
        f"- Q7 package map: `synthesis/ksnk_q7_goi_quy_trinh.md` ({len(packages)} packages)\n"
        f"- PDF mains remain scan-only (no local OCR engine / no Mistral key in env)\n"
    )
    log_path.write_text(log, encoding="utf-8")

    print(
        json.dumps(
            {
                "docx": len(all_att),
                "docx_ok": sum(1 for a in all_att if a["has_text"]),
                "enriched_concepts": enriched,
                "q7_packages": len(packages),
            },
            ensure_ascii=False,
        ),
        flush=True,
    )


if __name__ == "__main__":
    main()
