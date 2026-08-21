from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path
import fitz

out = Path(r'C:\Users\tantt\.nanobot\workspace\Phan_hoi_KSNK_phong_do_thinh_luc_20260722.docx')
r14o_pdf = Path(r'C:\Users\tantt\.nanobot\workspace\wiki\raw\manuals\Resonance_R14O_user_handbook.pdf')
rion_pdf = Path(r'C:\Users\tantt\.nanobot\workspace\wiki\raw\manuals\Rion_AA-M1C1_instruction_manual.pdf')
rs_pdf = Path(r'C:\Users\tantt\Downloads\New folder (4)\RS-H1_English Instruction manual_63290 (1).pdf')
img_dir = out.parent / '_manual_figures'
img_dir.mkdir(exist_ok=True)

def add_ocr_evidence(ocr_dir, title, pages, prefix):
    doc.add_paragraph(title, style='Caption')
    for page_no, printed_page in pages:
        source = Path(ocr_dir) / 'images' / f'page-{page_no:03d}.png'
        img = img_dir / f'{prefix}_p{printed_page}.png'
        from shutil import copyfile
        copyfile(source, img)
        doc.add_picture(str(img), width=Cm(15.5))
        cap = doc.add_paragraph(f'{title} – trang {printed_page}.')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2); sec.bottom_margin = Cm(2); sec.left_margin = Cm(2.5); sec.right_margin = Cm(2)
for name in ('Normal','Title','Heading 1','Heading 2'):
    st = doc.styles[name]
    st.font.name = 'Arial'; st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Heading 1'].font.size = Pt(14); doc.styles['Heading 1'].font.bold = True
doc.styles['Heading 2'].font.size = Pt(12); doc.styles['Heading 2'].font.bold = True

def heading(text, level=1): doc.add_heading(text, level=level)
def bullets(items):
    for item in items: doc.add_paragraph(item, style='List Bullet')

t = doc.add_paragraph(style='Title'); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.add_run('PHẢN HỒI KSNK\nPHÒNG ĐO THÍNH LỰC').bold = True
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Về vệ sinh và khử khuẩn trang thiết bị y tế\nNgày 22/07/2026').bold = True
heading('1. Mục đích')
doc.add_paragraph('Phản hồi các nội dung KSNK liên quan đến vệ sinh, khử khuẩn và bảo dưỡng thiết bị tại phòng đo thính lực, trên cơ sở quy trình KSNK và hướng dẫn sử dụng thiết bị đã thu thập.')
heading('2. Tài liệu đã rà soát')
bullets(['TA5.KSNK.QT.05 – Quy trình vệ sinh trang thiết bị y tế, Quận 7.','TA2.KSNK.QT.05 – Quy trình vệ sinh trang thiết bị, Tân Bình.','Hướng dẫn sử dụng thiết bị OAE Resonance/R14O, bản User Handbook Manual.','Instruction Manual Rion AA-M1C1 Audiometer.','Hướng dẫn sử dụng máy đo nhĩ lượng RS-H1.'])
heading('3. Nội dung phản hồi')
heading('3.1. Quy định chung đã xác minh', 2)
bullets(['Vệ sinh/khử khuẩn TTBYT trước và ngay sau khi sử dụng cho người bệnh theo phân loại và trách nhiệm của đơn vị.','Pha và sử dụng hóa chất theo hướng dẫn nhà sản xuất; không tự suy rộng hóa chất giữa các model.','Làm sạch từ sạch đến bẩn, từ trên xuống dưới; thay mặt khăn khi chuyển vị trí.','Sau thời gian tiếp xúc hóa chất, lau lại khi cần để loại bỏ tồn dư; kiểm tra tình trạng và hoạt động thiết bị.'])
heading('3.2. Thiết bị OAE Daewon', 2)
bullets(['Tắt thiết bị và rút phích cắm trước khi vệ sinh; không để chất lỏng lọt vào thiết bị hoặc phụ kiện. (Tài liệu hướng dẫn vệ sinh đầu dò OAE Daewon do đơn vị cung cấp).'])
doc.add_paragraph('Quy trình vệ sinh chuyên sâu đầu dò OAE Daewon:', style='List Bullet')
for i, item in enumerate(['Tháo lắp bên ngoài: mở rộng nhẹ hai ngàm nhựa ở hai bên đầu dò và kéo phần nhựa trong suốt ra ngoài.','Vệ sinh kênh dẫn: sử dụng đúng dụng cụ làm sạch đi kèm theo máy để làm sạch cả ba kênh trên đầu dò.','Kiểm tra trực quan: soi kỹ ba kênh dưới ánh sáng, bảo đảm không còn ráy tai hoặc mảnh vụn.','Lắp ráp lại: đặt phần nhựa trong suốt về phía sau; căn chỉnh phần lõm trên đầu khớp đúng với răng tương ứng trên thân trước.','Khóa ngàm: đẩy phần trước đầu dò về vị trí ban đầu; kiểm tra phần nhựa trong suốt khớp hoàn toàn và hai ngàm nhựa đã khóa. Không cố đóng đầu dò khi các chi tiết chưa khớp.'], 1):
    doc.add_paragraph(f'Bước {i}: {item}')
doc.add_paragraph('Núm tai dùng một lần phải thay trước người bệnh mới. Hóa chất vệ sinh thân máy và phụ kiện khác phải đối chiếu IFU đúng model; không suy rộng từ manual của thiết bị khác.')
heading('3.3. Thiết bị OAE Resonance/R14O', 2)
bullets(['Tắt thiết bị và rút phích cắm trước khi vệ sinh; không để chất lỏng lọt vào thiết bị hoặc phụ kiện. (R14O User Handbook Manual, trang 83).','Không dùng alcohol hoặc spirits; lau vỏ bằng khăn ẩm, không dùng chất tẩy rửa mạnh và tránh chất lỏng lọt vào trong. (R14O User Handbook Manual, trang 83).','Các bộ phận tiếp xúc người bệnh phải được làm sạch sau mỗi lần sử dụng bằng dung dịch sát khuẩn phù hợp theo mục bảo trì và vệ sinh của manual. Núm tai đầu dò là loại dùng một lần, không vô khuẩn, phải thay trước người bệnh mới. (R14O User Handbook Manual, trang 7, 83).','Trước khi sử dụng, kiểm tra hiệu năng đầu dò trong khoang kiểm tra chuyên dụng. Tháo phần đầu dò, lấy phần nhựa trong suốt và làm sạch 3 kênh bằng dụng cụ đi kèm; kiểm tra không còn mảnh vụn rồi lắp lại đúng khớp. (R14O User Handbook Manual, trang 73–74, 84–85).','Hiệu chuẩn, sửa chữa, cài đặt, cập nhật và thay đầu dò chỉ do nhân viên kỹ thuật đủ năng lực, được nhà sản xuất ủy quyền; R14O yêu cầu kiểm tra và hiệu chuẩn ít nhất mỗi năm. (R14O User Handbook Manual, trang 8, 83, 86).'])
doc.add_paragraph('Hình dẫn chứng từ manual R14O – mục vệ sinh và bảo trì đầu dò (trang in 83–85):', style='Caption')
_pdf = fitz.open(r14o_pdf)
for _page_no, _name in [(47, 'r14o_cleaning_p83.png'), (48, 'r14o_probe_p84.png')]:
    _pix = _pdf[_page_no].get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
    _img = img_dir / _name
    _pix.save(_img)
    doc.add_picture(str(_img), width=Cm(15.5))
    _cap = doc.add_paragraph(f'R14O User Handbook Manual, trang {_page_no + 36}.')
    _cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
_pdf.close()
heading('3.3. Audiometer Rion AA-M1C1', 2)
add_ocr_evidence(r'C:\Users\tantt\.nanobot\workspace\_ocr_verify\rion', 'Instruction Manual Rion AA-M1C1 – vệ sinh màn hình và thân máy', [(18, 16)], 'rion_clean')
doc.add_paragraph('Quy trình vệ sinh sau mỗi lần đo và kiểm tra định kỳ:')
for i, item in enumerate(['Tắt máy và rút nguồn trước khi vệ sinh.','Dùng bông y tế làm ẩm với cồn khử khuẩn để làm sạch tai nghe, đệm tai cao su, vòng đeo đầu và nút bấm phản hồi trước và sau mỗi lần đo.','Lau vỏ máy bằng cồn khử khuẩn hoặc chất tẩy rửa trung tính pha loãng, sau đó lau lại bằng nước sạch.','Lau màn hình LCD/bảng điều khiển cảm ứng chỉ bằng chất tẩy rửa trung tính pha loãng; không dùng cồn, dung môi pha loãng hoặc benzen trên màn hình.','Kiểm tra hằng ngày bằng mắt và nghe; theo dõi hiệu chuẩn chủ quan ít nhất mỗi tuần và hiệu chuẩn khách quan ít nhất mỗi năm.','Ngừng sử dụng vòng đeo đầu nếu lỏng, nứt hoặc hư hỏng; chuyển nhân viên kỹ thuật đủ năng lực đánh giá.'], 1):
    doc.add_paragraph(f'Bước {i}: {item}')
heading('3.4. Máy nhĩ lượng RS-H1', 2)
add_ocr_evidence(r'C:\Users\tantt\.nanobot\workspace\_ocr_verify\rs-h1', 'Instruction Manual RS-H1 – vệ sinh đầu dò và phụ kiện', [(18, 18)], 'rs_h1_clean')
doc.add_paragraph('Quy trình kiểm tra và vệ sinh trước/sau mỗi lần đo:')
for i, item in enumerate(['Kiểm tra trực quan và nghe máy; kiểm tra đầu dò tai, nút tai, dây dẫn và các bộ phận tiếp xúc người bệnh.','Lau đầu dò tai bằng khăn/gạc tẩm cồn khử khuẩn trước và sau mỗi lần đo; không để dung dịch lọt vào đầu dò và để khô hoàn toàn.','Xử lý nút tai riêng: rửa bằng dung dịch rửa pha loãng hoặc xà phòng phẫu thuật, tráng sạch, làm khô hoàn toàn, sau đó khử khuẩn giữa từng người bệnh. Không áp dụng hướng dẫn lau cồn của đầu dò tai cho nút tai.','Lau vòng đeo đầu và đệm tai cao su của tai nghe đường khí bằng khăn/gạc tẩm cồn khử khuẩn trước và sau mỗi lần đo.','Lau thân máy bằng khăn vắt kỹ với cồn khử khuẩn hoặc chất tẩy rửa trung tính pha loãng theo IFU; lau lại bằng khăn ẩm sạch. Không dùng thinner hoặc benzene.','Theo dõi hiệu chuẩn khách quan ít nhất mỗi năm một lần; sửa chữa và hiệu chuẩn do nhân viên kỹ thuật đủ năng lực thực hiện.'], 1):
    doc.add_paragraph(f'Bước {i}: {item}')
heading('4. Hành động khắc phục/duy trì đề xuất')
for i, item in enumerate(['Bổ sung bảng kiểm vệ sinh phòng đo thính lực, gồm trước đo, sau đo, đệm tai/vòng đeo đầu/nút bấm, đầu dò OAE, vỏ máy và LCD.','Dán hướng dẫn: không phun trực tiếp hóa chất; không để dung dịch lọt vào thiết bị; phân biệt hóa chất theo từng model.','Quản lý riêng núm tai nghe chèn dùng một lần và thay trước mỗi bệnh nhân.','Ghi nhận kiểm tra hằng ngày; theo dõi hiệu chuẩn chủ quan hằng tuần và khách quan hằng năm đối với AA-M1C1.','Đối chiếu model thực tế: nguồn trước ghi AA-MAC1, manual mới ghi AA-M1C1. Chỉ ban hành sau khi xác nhận nhãn máy, model và serial.'], 1):
    doc.add_paragraph(f'{i}. {item}')
heading('5. Điểm cần xác minh trước khi ban hành chính thức')
bullets(['Xác nhận model audiometer là AA-M1C1 hay AA-MAC1.','Lấy đúng IFU của OAE, máy nhĩ lượng và phụ kiện đang sử dụng.','Phòng TTBYT/nhà sản xuất xác nhận hóa chất tương thích cho từng bề mặt và phụ kiện.','Không áp dụng Cidex OPA hoặc thời gian ngâm 5 phút cho núm cao su nếu chưa có xác nhận từ IFU đúng model.'])
heading('6. Kết luận')
doc.add_paragraph('Các yêu cầu trên đủ cơ sở để xây dựng bảng kiểm và hướng dẫn thao tác tại phòng đo thính lực. Hóa chất và phương pháp xử lý phụ kiện phải được đối chiếu theo đúng model/IFU trước khi ban hành; hiệu chuẩn, sửa chữa và thay thế phải do nhân sự kỹ thuật đủ năng lực thực hiện.')
p = doc.add_paragraph(); p.add_run('Tài liệu tham chiếu nội bộ: ').bold = True; p.add_run('wiki/synthesis/phong_do_thinh_luc_ve_sinh_ttb_20260722.md')
for section in doc.sections:
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; footer.add_run('Phản hồi KSNK – Phòng đo thính lực | 22/07/2026')
doc.save(out)
print(out)
