# -*- coding: utf-8 -*-
"""Build KSNK wiki from Downloads New folder (3). Extract text + pages; no full binary copy."""
from __future__ import annotations

import hashlib
import json
import re
import shutil
from collections import defaultdict
from datetime import date
from pathlib import Path

from pypdf import PdfReader
from docx import Document

ROOT = Path.home() / "Downloads" / "New folder (3)"
WS = Path.home() / ".nanobot" / "workspace"
WIKI = WS / "wiki"
RAW = WIKI / "raw" / "ksnk"
EXTRACT = RAW / "extracts"
CONCEPTS = WIKI / "concepts"
SYN = WIKI / "synthesis"
ENT = WIKI / "entities"
TODAY = date.today().isoformat()

# Prefer Q7 QTVH main PDFs (operational). Skip QTKT scan duplicates when QTVH exists.
# Max pages to extract per PDF (header + body enough for wiki summary)
MAX_PAGES = 12
MAX_CHARS = 18000


def slugify(s: str) -> str:
    s = s.lower().strip()
    s = s.replace("đ", "d").replace("Đ", "d")
    # remove diacritics roughly via NFKD
    import unicodedata

    s = unicodedata.normalize("NFKD", s)
    s = "".join(c for c in s if not unicodedata.combining(c))
    s = re.sub(r"[^a-z0-9]+", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return s[:80]


def site_of(rel: Path) -> str:
    parts = rel.parts
    if not parts:
        return "unknown"
    p0 = parts[0]
    if "Quận" in p0 or p0.startswith("Quan"):
        return "q7"
    if "Tân" in p0 or "Tan" in p0:
        return "tb"
    return slugify(p0)


def doc_code(name: str) -> str | None:
    # TA5.KSNK.QT.02, TA2.KSNK.QĐ.03, KSNK.QT.01, CS.KSNK.02, QTKT.KSNK.01, QT.KSNK.09
    patterns = [
        r"(TA[25]\.KSNK\.(?:QT|QD|QĐ|HD)\.\d+[A-Z]?)",
        r"(KSNK\.(?:QT|QD|QĐ|HD)\.\d+[A-Z]?)",
        r"(CS\.KSNK\.\d+)",
        r"(QTKT\.KSNK\.\d+)",
        r"(QT\.KSNK\.\d+)",
        r"(TA5\.KSNK\.(?:QT|QD|QĐ|HD)\.\d+[A-Z]?)",
    ]
    # normalize Đ
    n = name.replace("QĐ", "QD").replace("qđ", "qd")
    for pat in patterns:
        m = re.search(pat, n, re.I)
        if m:
            code = m.group(1).upper().replace("QD", "QĐ")
            # fix TA5.KSNK.QD -> QĐ already
            return code
    return None


def title_from_name(name: str) -> str:
    stem = Path(name).stem
    # drop leading code_
    stem = re.sub(
        r"^(TA[25]\.|)?KSNK\.(QT|QĐ|QD|HD)\.\d+[A-Z]?[_ ]*",
        "",
        stem,
        flags=re.I,
    )
    stem = re.sub(r"^(CS\.KSNK\.\d+|QTKT\.KSNK\.\d+|QT\.KSNK\.\d+)[_ ]*", "", stem, flags=re.I)
    stem = re.sub(r"_+", " ", stem).strip(" -_")
    stem = re.sub(r"\s+", " ", stem)
    return stem or Path(name).stem


def extract_pdf(path: Path) -> str:
    try:
        reader = PdfReader(str(path))
        parts = []
        n = min(len(reader.pages), MAX_PAGES)
        for i in range(n):
            try:
                t = reader.pages[i].extract_text() or ""
            except Exception:
                t = ""
            if t.strip():
                parts.append(f"--- page {i+1} ---\n{t.strip()}")
        text = "\n\n".join(parts)
        if len(text) > MAX_CHARS:
            text = text[:MAX_CHARS] + "\n\n[...truncated...]"
        return text
    except Exception as e:
        return f"[extract_error] {e}"


def extract_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
        paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        # tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                line = " | ".join(cells)
                if line.strip(" |"):
                    paras.append(line)
        text = "\n".join(paras)
        if len(text) > MAX_CHARS:
            text = text[:MAX_CHARS] + "\n\n[...truncated...]"
        return text
    except Exception as e:
        return f"[extract_error] {e}"


def is_main_doc(path: Path) -> bool:
    """Main process PDF (not BK/BM/PL attachment alone)."""
    name = path.name
    if name.startswith("~$") or name.lower() == "thumbs.db":
        return False
    if path.suffix.lower() != ".pdf":
        return False
    # attachments often start with BK/BM/PL
    if re.match(r"^(BK|BM|PL|QA)\d*", name, re.I):
        return False
    if re.match(r"^PL\.\d+", name, re.I):
        return False
    return True


def pick_preferred_mains(files: list[Path]) -> list[Path]:
    """One preferred main PDF per (site, code). Prefer QTVH over QTKT; newer dated name."""
    groups: dict[tuple[str, str], list[Path]] = defaultdict(list)
    no_code = []
    for p in files:
        if not is_main_doc(p):
            continue
        rel = p.relative_to(ROOT)
        site = site_of(rel)
        code = doc_code(p.name) or doc_code(str(rel))
        if not code:
            no_code.append(p)
            continue
        groups[(site, code)].append(p)

    chosen = []
    for key, paths in groups.items():
        def score(p: Path):
            s = str(p).lower()
            sc = 0
            if "qtvh" in s:
                sc += 100
            if "qtkt" in s:
                sc -= 50
            if "chương trình" in s or "chuong trinh" in s:
                sc += 20
            # prefer longer path with folder = more complete package? prefer larger file slightly
            sc += min(p.stat().st_size / 1e6, 30)  # up to +30
            # dated version
            if re.search(r"\d{2}\.\d{2}\.\d{4}", p.name):
                sc += 5
            return sc

        paths.sort(key=score, reverse=True)
        chosen.append(paths[0])
    # include no_code mains that look like process docs
    for p in no_code:
        chosen.append(p)
    return sorted(chosen, key=lambda p: str(p).lower())


def summarize_from_text(text: str, title: str) -> dict:
    """Heuristic extract of purpose / steps / notes from Vietnamese SOPs."""
    lines = [ln.strip() for ln in text.splitlines() if ln.strip() and not ln.startswith("--- page")]
    # drop very short noise
    body = "\n".join(lines)
    purpose = ""
    for key in ["1. Mục đích", "I. Mục đích", "Mục đích", "1. MỤC ĐÍCH", "MỤC ĐÍCH"]:
        if key in body:
            idx = body.find(key)
            chunk = body[idx : idx + 500]
            # next section
            m = re.search(
                r"(?:Mục đích|MỤC ĐÍCH)[:\s]*(.+?)(?:\n\s*(?:\d+\.|II\.|2\.|Phạm vi|PHẠM VI|Đối tượng)|$)",
                chunk,
                re.S | re.I,
            )
            if m:
                purpose = re.sub(r"\s+", " ", m.group(1)).strip(" :.-")[:400]
            break
    scope = ""
    for key in ["2. Phạm vi", "II. Phạm vi", "Phạm vi áp dụng", "PHẠM VI"]:
        if key.lower() in body.lower():
            idx = body.lower().find(key.lower())
            chunk = body[idx : idx + 400]
            m = re.search(
                r"(?:Phạm vi(?: áp dụng)?|PHẠM VI)[:\s]*(.+?)(?:\n\s*(?:\d+\.|III\.|3\.|Tài liệu|Định nghĩa)|$)",
                chunk,
                re.S | re.I,
            )
            if m:
                scope = re.sub(r"\s+", " ", m.group(1)).strip(" :.-")[:400]
            break

    # collect numbered steps-ish lines
    steps = []
    for ln in lines:
        if re.match(r"^(\d+[\.\)]\s+|[-•]\s+|[a-z][\.\)]\s+)", ln) and len(ln) > 15:
            steps.append(ln[:200])
        if len(steps) >= 12:
            break

    # regulatory refs
    refs = sorted(
        set(
            re.findall(
                r"(?:Thông tư|Nghị định|Quyết định|QĐ|TT|NĐ|Luật|ISO|CDC|WHO)[^\n,]{0,80}",
                body,
                flags=re.I,
            )
        )
    )[:15]

    return {
        "purpose": purpose,
        "scope": scope,
        "steps": steps,
        "refs": refs,
        "char_count": len(text),
        "has_text": len(text.strip()) > 80 and not text.startswith("[extract_error]"),
    }


def md_escape(s: str) -> str:
    return s.replace("|", "\\|")


def build():
    RAW.mkdir(parents=True, exist_ok=True)
    EXTRACT.mkdir(parents=True, exist_ok=True)
    CONCEPTS.mkdir(parents=True, exist_ok=True)
    SYN.mkdir(parents=True, exist_ok=True)

    all_files = [
        p
        for p in ROOT.rglob("*")
        if p.is_file() and not p.name.startswith("~$") and p.name.lower() != "thumbs.db"
    ]
    print(f"all_files={len(all_files)}", flush=True)

    # catalog
    catalog_lines = [
        "---",
        "type: source",
        'title: "Catalog nguồn KSNK (Quận 7 + Tân Bình)"',
        "status: draft",
        f"updated: {TODAY}",
        "tags: [ksnk, infection-control, catalog]",
        "---",
        "",
        f"# Catalog nguồn Kiểm soát nhiễm khuẩn",
        "",
        f"- Nguồn gốc (immutable path): `{ROOT}`",
        f"- Ngày ingest: {TODAY}",
        f"- Tổng file (bỏ `~$`, Thumbs): **{len(all_files)}**",
        "- Không copy binary full vào wiki (dung lượng lớn); giữ path + extract text các SOP chính.",
        "",
        "## Danh mục",
        "",
        "| Site | Loại | Mã (ước) | File | KB |",
        "|---|---|---|---|---|",
    ]
    by_site = defaultdict(int)
    for p in sorted(all_files, key=lambda x: str(x).lower()):
        rel = p.relative_to(ROOT)
        site = site_of(rel)
        by_site[site] += 1
        code = doc_code(p.name) or doc_code(str(rel)) or ""
        kind = p.suffix.lower().lstrip(".")
        catalog_lines.append(
            f"| {site} | {kind} | {code} | `{rel.as_posix()}` | {p.stat().st_size//1024} |"
        )
    catalog_path = RAW / "catalog.md"
    catalog_path.write_text("\n".join(catalog_lines) + "\n", encoding="utf-8")
    print(f"catalog -> {catalog_path} sites={dict(by_site)}", flush=True)

    mains = pick_preferred_mains(all_files)
    print(f"main_pdfs={len(mains)}", flush=True)

    records = []
    for p in mains:
        rel = p.relative_to(ROOT)
        site = site_of(rel)
        code = doc_code(p.name) or doc_code(str(rel)) or f"NOCODE_{slugify(p.stem)[:40]}"
        title = title_from_name(p.name)
        print(f"extract {site} {code.encode('ascii','replace').decode()} {p.stat().st_size//1024}KB ...", flush=True)
        text = extract_pdf(p)
        meta = summarize_from_text(text, title)
        # save extract
        ex_name = f"{site}_{slugify(code)}_{slugify(title)[:40]}.md"
        ex_path = EXTRACT / ex_name
        ex_body = [
            "---",
            "type: source",
            f'title: "{code} — {title}"',
            "status: draft",
            f"updated: {TODAY}",
            f'sources: ["{rel.as_posix()}"]',
            "tags: [ksnk, extract]",
            "---",
            "",
            f"# {code} — {title}",
            "",
            f"- Site: **{site}**",
            f"- Source path: `{p}`",
            f"- Relative: `{rel.as_posix()}`",
            f"- Size: {p.stat().st_size//1024} KB",
            f"- Extract pages: ≤{MAX_PAGES}",
            f"- Text layer: {'ok' if meta['has_text'] else 'weak/scan'}",
            "",
            "## Extracted text",
            "",
            "```",
            text if text.strip() else "[empty]",
            "```",
            "",
        ]
        ex_path.write_text("\n".join(ex_body), encoding="utf-8")
        rec = {
            "site": site,
            "code": code,
            "title": title,
            "rel": rel.as_posix(),
            "abs": str(p),
            "size_kb": p.stat().st_size // 1024,
            "extract": f"raw/ksnk/extracts/{ex_name}",
            "meta": meta,
        }
        records.append(rec)

    (RAW / "mains_index.json").write_text(
        json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # Group by normalized process key (code without site prefix)
    def norm_key(code: str) -> str:
        c = code.upper().replace("QĐ", "QD")
        c = re.sub(r"^TA[25]\.", "", c)
        c = c.replace("QD", "QĐ")
        return c

    groups: dict[str, list[dict]] = defaultdict(list)
    for r in records:
        groups[norm_key(r["code"])].append(r)

    # Topic pages for each process group
    concept_pages = []
    for key, items in sorted(groups.items(), key=lambda kv: kv[0]):
        # prefer q7 item for primary content
        primary = next((i for i in items if i["site"] == "q7"), items[0])
        slug = slugify(f"ksnk_{key}_{primary['title']}")
        page_path = CONCEPTS / f"{slug}.md"
        # if collision, add site
        if page_path.exists() and key in page_path.read_text(encoding="utf-8"):
            pass

        sites = ", ".join(sorted({i["site"] for i in items}))
        sources = []
        for i in items:
            sources.append(f'  - "{i["extract"]}"')
            sources.append(f'  - "source-path:{i["rel"]}"')

        meta = primary["meta"]
        lines = [
            "---",
            "type: concept",
            f'title: "{primary["code"]} — {primary["title"]}"',
            "status: draft",
            "sources:",
            *sources,
            f"updated: {TODAY}",
            "tags: [ksnk, infection-control, quy-trinh]",
            "refs: []",
            "---",
            "",
            f"# {primary['code']} — {primary['title']}",
            "",
            f"- Nhóm mã: `{key}`",
            f"- Cơ sở có tài liệu: **{sites}**",
            f"- Trang tổng hợp: [[synthesis/ksnk_quy_trinh_hub]]",
            "",
            "## Tóm tắt (từ extract)",
            "",
        ]
        if meta.get("purpose"):
            lines += [f"**Mục đích (trích):** {meta['purpose']}", ""]
        else:
            lines += [
                "*Chưa trích được mục đích rõ (PDF scan/OCR yếu hoặc layout phức tạp). Xem extract thô.*",
                "",
            ]
        if meta.get("scope"):
            lines += [f"**Phạm vi (trích):** {meta['scope']}", ""]

        lines += ["## Biến thể theo cơ sở", ""]
        for i in sorted(items, key=lambda x: x["site"]):
            flag = "text-ok" if i["meta"]["has_text"] else "text-weak"
            lines.append(
                f"- **{i['site'].upper()}** `{i['code']}` — {i['title']} ({i['size_kb']} KB, {flag})"
            )
            lines.append(f"  - Nguồn: `{i['rel']}`")
            lines.append(f"  - Extract: [[{i['extract'].replace('.md','')}]]".replace("raw/ksnk/extracts/", "raw path extracts/ "))
            lines.append(f"  - File extract: `{i['extract']}`")

        if meta.get("steps"):
            lines += ["", "## Gợi ý các bước / mục (heuristic từ text)", ""]
            for st in meta["steps"][:10]:
                lines.append(f"- {st}")

        if meta.get("refs"):
            lines += ["", "## Tài liệu / căn cứ được nhắc trong extract", ""]
            for rf in meta["refs"]:
                lines.append(f"- {rf}")

        lines += [
            "",
            "## Ghi chú",
            "",
            "- Nội dung wiki **không thay** SOP gốc; chỉ index + tóm tắt có citation.",
            "- PDF scan: cần OCR đầy đủ nếu text-layer yếu.",
            "- Không suy diễn yêu cầu quy định ngoài những gì có trong nguồn.",
            "",
        ]
        page_path.write_text("\n".join(lines), encoding="utf-8")
        concept_pages.append(
            {
                "slug": slug,
                "path": f"concepts/{slug}",
                "title": f"{primary['code']} — {primary['title']}",
                "key": key,
                "sites": sites,
                "has_text": any(i["meta"]["has_text"] for i in items),
            }
        )
        print(f"page {page_path.name}", flush=True)

    # Hub synthesis
    hub = SYN / "ksnk_quy_trinh_hub.md"
    hub_lines = [
        "---",
        "type: synthesis",
        'title: "Hub quy trình Kiểm soát nhiễm khuẩn (KSNK)"',
        "status: draft",
        "sources:",
        '  - "raw/ksnk/catalog.md"',
        '  - "raw/ksnk/mains_index.json"',
        f'  - "source-root:{ROOT.as_posix()}"',
        f"updated: {TODAY}",
        "tags: [ksnk, infection-control, hub]",
        "refs: []",
        "---",
        "",
        "# Hub quy trình Kiểm soát nhiễm khuẩn (KSNK)",
        "",
        f"Ingest {TODAY} từ `{ROOT}`.",
        "",
        "## Phạm vi nguồn",
        "",
        f"- **Quận 7 (TA5)** + **Tân Bình (TA2)**: thư mục `20. KIỂM SOÁT NHIỄM KHUẨN`.",
        f"- Tổng file catalog: **{len(all_files)}** (pdf/docx/xlsx…).",
        f"- SOP chính (PDF main, dedupe theo mã): **{len(records)}** → **{len(concept_pages)}** trang concept.",
        "- Binary gốc **không** copy full vào wiki (dung lượng lớn). Provenance = absolute path + catalog.",
        "",
        "## Cấu trúc mã tài liệu",
        "",
        "- `TA5.KSNK.*` — PK/BV Tâm Anh Quận 7",
        "- `TA2.KSNK.*` / `KSNK.*` — Tân Bình",
        "- Tiền tố loại: `QT` quy trình, `QĐ` quy định, `HD` hướng dẫn, `CS` chương trình, `QTKT` quy trình kỹ thuật",
        "- Kèm theo: `BK` bảng kiểm, `BM` biểu mẫu, `PL` phụ lục (liệt kê trong catalog, chưa tách page riêng)",
        "",
        "## Danh mục quy trình (concept pages)",
        "",
        "| Mã nhóm | Trang | Cơ sở | Text |",
        "|---|---|---|---|",
    ]
    for cp in sorted(concept_pages, key=lambda x: x["key"]):
        hub_lines.append(
            f"| `{cp['key']}` | [[{cp['path']}|{cp['title']}]] | {cp['sites']} | {'ok' if cp['has_text'] else 'weak'} |"
        )

    # Q7-focused quick list
    q7 = [r for r in records if r["site"] == "q7"]
    hub_lines += [
        "",
        "## Ưu tiên Quận 7 (TA5) — SOP main",
        "",
    ]
    for r in sorted(q7, key=lambda x: x["code"]):
        hub_lines.append(
            f"- `{r['code']}` — {r['title']} — `{r['rel']}` — extract `{r['extract']}`"
        )

    weak = [r for r in records if not r["meta"]["has_text"]]
    hub_lines += [
        "",
        "## Text-layer yếu / cần OCR",
        "",
    ]
    if not weak:
        hub_lines.append("- (không có trong lô main đã extract)")
    else:
        for r in weak:
            hub_lines.append(f"- `{r['code']}` {r['site']} — `{r['rel']}`")

    hub_lines += [
        "",
        "## Việc còn lại",
        "",
        "- OCR đầy đủ PDF scan (Mistral/doc-ocr) cho mục text-weak.",
        "- Tách page riêng cho BK/BM quan trọng (giám sát VST, chất thải, CSSD).",
        "- Map RACI / liên kết MEIMS modules nếu đưa vào phần mềm.",
        "- Đối chiếu chéo TA5 vs TA2 khi cùng mã QT (khác phạm vi phòng khám vs bệnh viện).",
        "",
        "## An toàn",
        "",
        "- Wiki là **tra cứu nội bộ**, không thay thế văn bản kiểm soát tài liệu đang hiệu lực tại cơ sở.",
        "- Không suy diễn yêu cầu BYT ngoài citation trong SOP.",
        "",
    ]
    hub.write_text("\n".join(hub_lines), encoding="utf-8")
    print(f"hub -> {hub}", flush=True)

    # Entity stub for KSNK function
    ent = ENT / "ksnk_tam_anh.md"
    ent.write_text(
        "\n".join(
            [
                "---",
                "type: entity",
                'title: "Kiểm soát nhiễm khuẩn — Hệ thống Tâm Anh (TA2/TA5)"',
                "status: draft",
                "sources:",
                '  - "raw/ksnk/catalog.md"',
                '  - "synthesis/ksnk_quy_trinh_hub.md"',
                f"updated: {TODAY}",
                "tags: [ksnk, organization]",
                "refs: []",
                "---",
                "",
                "# Kiểm soát nhiễm khuẩn — Hệ thống Tâm Anh",
                "",
                "- **TA5**: PK Đa khoa Tâm Anh Quận 7 — mã tài liệu `TA5.KSNK.*`",
                "- **TA2**: BV/cơ sở Tân Bình — mã `TA2.KSNK.*` / `KSNK.*`",
                "- Hub quy trình: [[synthesis/ksnk_quy_trinh_hub]]",
                "- Liên quan: [[entities/tam-anh-clinic-q7]]",
                "",
                "Chức năng: xây dựng/ban hành QT–QĐ–HD KSNK; giám sát tuân thủ; CSSD/đồ vải/chất thải/VST/PPE tùy cơ sở.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    # Update index.md
    index_path = WIKI / "index.md"
    idx = index_path.read_text(encoding="utf-8")
    block_topics = ""
    if "ksnk_quy_trinh_hub" not in idx:
        # append sections
        add = [
            "",
            "## KSNK — Kiểm soát nhiễm khuẩn",
            f"- [[synthesis/ksnk_quy_trinh_hub]] — Hub quy trình KSNK (TA5 Q7 + TA2 Tân Bình), ingest {TODAY}",
            f"- [[entities/ksnk_tam_anh]] — Đơn vị/chức năng KSNK Tâm Anh",
            f"- [[raw/ksnk/catalog]] — Catalog {len(all_files)} file nguồn",
            "",
            "### Quy trình KSNK (concepts)",
        ]
        for cp in sorted(concept_pages, key=lambda x: x["key"]):
            add.append(f"- [[{cp['path']}]] — {cp['title']} ({cp['sites']})")
        idx = idx.rstrip() + "\n" + "\n".join(add) + "\n"
        index_path.write_text(idx, encoding="utf-8")
        print("index updated", flush=True)
    else:
        print("index already has hub", flush=True)

    # log
    log_path = WIKI / "log.md"
    log = log_path.read_text(encoding="utf-8") if log_path.exists() else "# Wiki Log\n\n"
    entry = (
        f"\n## {TODAY} — ingest KSNK\n"
        f"- Source: `{ROOT}`\n"
        f"- Catalog: `raw/ksnk/catalog.md` ({len(all_files)} files)\n"
        f"- Main SOP PDF extracted: {len(records)} → concepts: {len(concept_pages)}\n"
        f"- Hub: `synthesis/ksnk_quy_trinh_hub.md`\n"
        f"- Entity: `entities/ksnk_tam_anh.md`\n"
        f"- Note: no full binary copy; text extract ≤{MAX_PAGES} pages/PDF\n"
        f"- Weak text-layer: {len(weak)} docs\n"
    )
    log_path.write_text(log.rstrip() + "\n" + entry + "\n", encoding="utf-8")

    # cleanup helpers
    for h in [
        WS / "_ksnk_inventory.txt",
        WS / "_ksnk_build_wiki.py",
    ]:
        try:
            if h.exists():
                # keep script until end - delete inventory only here; script deletes self last
                if h.name.endswith(".txt"):
                    h.unlink()
        except OSError:
            pass

    print(
        json.dumps(
            {
                "files": len(all_files),
                "mains": len(records),
                "concepts": len(concept_pages),
                "weak": len(weak),
                "q7_mains": len(q7),
            },
            ensure_ascii=False,
        ),
        flush=True,
    )


if __name__ == "__main__":
    build()
